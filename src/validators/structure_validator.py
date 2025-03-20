class StructureValidator:
    """文档结构验证器基类"""
    
    def validate(self, doc_path: str) -> dict:
        """执行结构验证"""
        raise NotImplementedError("子类必须实现此方法")

class DocxStructureValidator(StructureValidator):
    """DOCX格式文档结构验证"""
    
    def validate(self, doc_path: str) -> dict:
        """执行核心验证流程"""
        doc = self._load_document(doc_path)
        return {
            "section_structure": self._validate_sections(doc),
            "reference_format": self._check_references(doc)
        }
    
    def _load_document(self, doc_path: str):
        from docx import Document
        return Document(doc_path)
    
    def _validate_sections(self, doc) -> dict:
        headings = self._extract_headings(doc)
        return {
            "required_sections": ["abstract", "introduction"],
            "found_sections": [h.text.lower() for h in headings],
            "missing_sections": self._find_missing_sections(headings)
        }
    
    def _extract_headings(self, doc):
        return [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
    
    def _find_missing_sections(self, headings):
        """改进中文章节模糊匹配"""
        section_mapping = {
            'introduction': ['引言', '研究背景', '背景与意义'],
            'conclusion': ['结论', '总结', '研究展望']
        }
        
        found_sections = set()
        for h in headings:
            text = h.text.strip().lower()
            for eng_name, zh_names in section_mapping.items():
                if any(zh in text for zh in zh_names):
                    found_sections.add(eng_name)
        
        return list(set(section_mapping.keys()) - found_sections)

    def _check_references(self, doc) -> dict:
        ref_section = self._find_reference_section(doc)
        return {
            "exists": bool(ref_section),
            "format_valid": self._validate_reference_format(ref_section) if ref_section else False
        }
    
    def _find_reference_section(self, doc):
        for para in doc.paragraphs:
            if "references" in para.text.lower():
                return para
        return None
    
    def _validate_reference_format(self, ref_section):
        return any(num in ref_section.text for num in ["[1]", "(1)"])

    def _check_fonts(self, doc) -> bool:
        """检查正文字体一致性"""
        fonts = set()
        for p in doc.paragraphs:
            if p.style.name == 'Normal':
                fonts.add(p.style.font.name)
        return len(fonts) == 1

    def _check_paragraph_spacing(self, doc) -> bool:
        """增强版段落间距检测"""
        spacings = set()
        for p in doc.paragraphs:
            if p.style.name == 'Normal':
                # 添加单位转换逻辑（Pt -> 磅）
                if p.paragraph_format.space_after:
                    spacings.add(round(p.paragraph_format.space_after.pt))
        return len(spacings) == 1

    def validate_style_consistency(self, doc_path: str) -> dict:
        """样式一致性验证"""
        doc = self._load_document(doc_path)
        return {
            "font_consistency": self._check_fonts(doc),
            "paragraph_spacing": self._check_paragraph_spacing(doc),
            "cover_structure": self._check_cover(doc)  # 添加封面检查
        }

    def _check_cover(self, doc):
        """验证封面必要元素（包含LOGO检测）"""
        # 添加对inline_shapes的类型检查
        logo_exists = False
        if hasattr(doc, 'inline_shapes'):
            logo_exists = len(doc.inline_shapes) > 0
            
        required = {
            'logo': logo_exists,  # 修改为安全访问方式
            'title': False,
            'school': False
        }
        
        for para in doc.paragraphs:
            text = para.text.replace('\t', '').replace(' ', '')
            if '题目：' in text and len(text) > 15:
                required['title'] = True
            elif '学院：' in text and '外国语' in text:
                required['school'] = True

        # 将字段键名转换为中文描述
        field_mapping = {
            'logo': '学校LOGO',
            'title': '论文标题',
            'school': '学院名称'
        }
                
        return {
            'logo_exists': required['logo'],
            'exists': all(required.values()),
            'missing_fields': [field_mapping[k] for k, v in required.items() if not v]
        }

    # 删除嵌套的DocxStructureValidator类定义
    # 原_analyze_toc方法修正到当前类中
    def _analyze_toc(self, doc):
        """分析目录结构"""
        toc_headings = [
            para for para in doc.paragraphs
            if para.style.name.startswith('Heading') and "目录" in para.text
        ]
        
        if not toc_headings:
            return {'valid': False, 'reason': '目录缺失'}
            
        return {
            'valid': True,
            'levels': len(set(p.style.name for p in doc.paragraphs 
                            if p.style.name.startswith('Heading')))
        }