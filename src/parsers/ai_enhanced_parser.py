# 新增AI解析基类
class AIParser:
    def __init__(self, config):
        self.ai_enabled = config['ai_integration']['deepseek_r1']['enabled']
        self.config = config['ai_integration']
        self._setup_http_client()  # 新增HTTP客户端初始化

    def _setup_http_client(self):
        """创建带重试机制的HTTP客户端"""
        import httpx
        self.client = httpx.Client(
            base_url=self.config['deepseek_r1']['api_base'],
            timeout=30.0,
            transport=httpx.HTTPTransport(retries=3)
        )

    def _call_deepseek_api(self, text, module):
        """实现API调用核心逻辑（新增完整实现）"""
        headers = {
            "Authorization": f"Bearer {self._load_api_key()}",
            "Content-Type": "application/json"
        }
        
        try:
            response = self.client.post(
                "/chat/completions",
                json={
                    "model": "deepseek-r1",
                    "messages": [{
                        "role": "user",
                        "content": f"请分析以下论文内容：{text[:2000]}"
                    }]
                },
                headers=headers
            )
            response.raise_for_status()
            return response.json()
        except httpx.HTTPStatusError as e:
            if e.response.status_code >= 500:
                # 服务器错误时重试
                raise RetryableError("Server error, retrying...")
            else:
                # 客户端错误停止重试
                raise PermanentError(f"API request failed: {e}")

    def _load_api_key(self):
        # 修正配置路径访问
        with open(self.config['deepseek_r1']['api_key_path']) as f:  # 修改路径访问方式
            return yaml.safe_load(f)['api_key']

# 在现有解析器中继承AI能力
class DocxParser(AIParser):
    def __init__(self, config):
        super().__init__(config)  # 必须调用父类初始化
        # 保留原有解析器成员变量
        self.doc = None  
        
    def parse_advanced(self, file_path):
        self.doc = Document(file_path)
        result = self._basic_parse()  # 原有解析逻辑
        
        if self.ai_enabled and self.config['deepseek_r1']['modules']['content_evaluation']:
            result['ai_analysis'] = self._enhance_with_ai()  # 新增AI增强分析
            
        return result

    def analyze_with_ai(self, text, module):  # 新增AI分析方法
        """执行具体的AI分析"""
        return self._call_deepseek_api(text, module)

    def _enhance_with_ai(self):
        """执行AI增强分析（新增方法）"""
        analysis = {}
        # 修正方法调用参数
        analysis['title_evaluation'] = self.analyze_with_ai(
            self._extract_title_text(),
            'content_evaluation'  # 修正模块名称匹配配置
        )
        return analysis

    def _basic_parse(self):  # 新增基础解析方法
        """模拟基础解析逻辑"""
        return {
            'title': self._extract_title_text(),
            'sections': [p.text for p in self.doc.paragraphs]
        }
    def _extract_title_text(self):  # 新增标题提取方法
        """从文档中提取标题文本"""
        if not self.doc.paragraphs:
            return ""
        return self.doc.paragraphs[0].text
from docx import Document  # 新增导入
import yaml
import httpx

# 在文件顶部新增自定义异常
class RetryableError(Exception):
    """可重试错误基类"""
    
class PermanentError(Exception):
    """不可恢复错误基类"""